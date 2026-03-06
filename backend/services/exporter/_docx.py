"""Generate ATS-friendly Word (.docx) documents from a CvVersion."""

from __future__ import annotations

import io
import uuid

from sqlalchemy.ext.asyncio import AsyncSession

from backend.models.tables import CvVersion

from ._context import _build_cv_context


async def generate_docx(db: AsyncSession, cv_version: CvVersion, user_id: uuid.UUID) -> bytes:
    """Generate a clean, ATS-friendly Word document (.docx) from the final CV data.

    Uses python-docx. No tables, headers, or footers — plain text formatted with
    paragraph styles for maximum ATS compatibility.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    context = await _build_cv_context(db, cv_version, user_id)
    profile = context["profile"]

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Default body font ────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def _add_section_heading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_bullet(text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        p.add_run(str(text)).font.size = Pt(10.5)

    def _add_item_header(left: str, right: str, subtitle: str = "", date_range: str = "") -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        run_left = p.add_run(left)
        run_left.bold = True
        run_left.font.size = Pt(11)
        if right:
            p.add_run("\t").font.size = Pt(10.5)
            p.add_run(right).font.size = Pt(10.5)
        if subtitle or date_range:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(0)
            parts = [x for x in [subtitle, date_range] if x]
            p2.add_run(" | ".join(parts)).font.size = Pt(10.5)

    # ── Name header ──────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(profile.get("name") or "Name")
    name_run.bold = True
    name_run.font.size = Pt(16)

    # ── Contact line ─────────────────────────────────────────────────────────
    contact_parts = [v for k in ["email", "phone", "linkedin_url", "location"] if (v := profile.get(k))]
    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(4)
        for run in contact_para.runs:
            run.font.size = Pt(10)

    # ── Education ────────────────────────────────────────────────────────────
    education = context.get("education", [])
    if education:
        _add_section_heading("Education")
        for edu in education:
            dates = " – ".join(filter(None, [edu.get("date_start"), edu.get("date_end")]))
            _add_item_header(
                left=edu.get("institution") or "",
                right=dates,
                subtitle=edu.get("degree") or "",
                date_range=edu.get("grade") or "",
            )
            for ach in edu.get("achievements") or []:
                _add_bullet(ach)
            mods = edu.get("modules") or []
            if mods:
                doc.add_paragraph(f"Relevant modules: {', '.join(mods)}").paragraph_format.left_indent = Inches(0.25)

    # ── Experience ───────────────────────────────────────────────────────────
    experiences = context.get("experiences", [])
    if experiences:
        _add_section_heading("Experience")
        for exp in experiences:
            dates = " – ".join(filter(None, [exp.get("date_start"), exp.get("date_end")]))
            _add_item_header(
                left=exp.get("company") or "",
                right=dates,
                subtitle=exp.get("role_title") or "",
                date_range=exp.get("location") or "",
            )
            for bullet in exp.get("bullets") or []:
                _add_bullet(bullet)

    # ── Projects ─────────────────────────────────────────────────────────────
    projects = context.get("projects", [])
    if projects:
        _add_section_heading("Projects")
        for proj in projects:
            dates = " – ".join(filter(None, [proj.get("date_start"), proj.get("date_end")]))
            _add_item_header(left=proj.get("name") or "", right=dates)
            if proj.get("description"):
                doc.add_paragraph(proj["description"]).paragraph_format.left_indent = Inches(0.25)
            for bullet in proj.get("bullets") or []:
                _add_bullet(bullet)

    # ── Leadership & Activities ───────────────────────────────────────────────
    activities = context.get("activities", [])
    if activities:
        _add_section_heading("Leadership & Activities")
        for act in activities:
            dates = " – ".join(filter(None, [act.get("date_start"), act.get("date_end")]))
            _add_item_header(
                left=act.get("organization") or "",
                right=dates,
                subtitle=act.get("role_title") or "",
            )
            for bullet in act.get("bullets") or []:
                _add_bullet(bullet)

    # ── Technical Skills ─────────────────────────────────────────────────────
    skills_by_category = context.get("skills_by_category", {})
    if skills_by_category:
        _add_section_heading("Technical Skills")
        for cat, skills in skills_by_category.items():
            if skills:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run_cat = p.add_run(f"{cat}: ")
                run_cat.bold = True
                run_cat.font.size = Pt(10.5)
                p.add_run(", ".join(skills)).font.size = Pt(10.5)

    # ── Serialise to bytes ───────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
